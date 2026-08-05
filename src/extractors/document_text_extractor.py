# -*- coding: utf-8 -*-

from io import BytesIO
from urllib.parse import urlparse
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader

from src.models.document_evidence import (
    DocumentPageEvidence,
    ParsedDocument,
)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ODT_MIME = "application/vnd.oasis.opendocument.text"
DOC_MIME = "application/msword"


# 依內容類型或網址副檔名判斷可解析文件格式。
def detect_document_kind(content_type: str, source_url: str) -> str:
    normalized_type = content_type.split(";", maxsplit=1)[0].strip().lower()
    suffix = urlparse(source_url).path.lower()
    if normalized_type == PDF_MIME or suffix.endswith(".pdf"):
        return "pdf"
    if normalized_type == DOCX_MIME or suffix.endswith(".docx"):
        return "docx"
    if normalized_type == ODT_MIME or suffix.endswith(".odt"):
        return "odt"
    if normalized_type == DOC_MIME or suffix.endswith(".doc"):
        return "doc_legacy"
    return "unsupported"


# 將支援文件轉成保留頁碼的結構化證據。
def extract_document(
    content: bytes,
    content_type: str,
    source_url: str,
    max_pdf_pages: int,
) -> ParsedDocument:
    kind = detect_document_kind(content_type, source_url)
    if kind == "pdf":
        return ParsedDocument(kind, _extract_pdf_pages(content, max_pdf_pages))
    if kind == "docx":
        return ParsedDocument(kind, _single_page(_extract_docx_text(content)))
    if kind == "odt":
        return ParsedDocument(kind, _single_page(_extract_odt_text(content)))
    if kind == "doc_legacy":
        raise ValueError("舊式 DOC 無法安全解析，需轉成 PDF、DOCX 或人工確認")
    raise ValueError("不支援的附件格式")


# 保留舊介面，讓既有資格規則仍可取得合併文字。
def extract_document_text(
    content: bytes,
    content_type: str,
    source_url: str,
    max_pdf_pages: int,
) -> str:
    return extract_document(content, content_type, source_url, max_pdf_pages).text


# 逐頁擷取 PDF；空白頁略過但原始頁碼不得重排。
def _extract_pdf_pages(
    content: bytes,
    max_pages: int,
) -> tuple[DocumentPageEvidence, ...]:
    reader = PdfReader(BytesIO(content))
    pages: list[DocumentPageEvidence] = []
    for page_number, page in enumerate(reader.pages[:max_pages], start=1):
        text = _normalize_text(page.extract_text() or "")
        if text:
            pages.append(DocumentPageEvidence(page_number, text))
    if not pages:
        raise ValueError("PDF 沒有可擷取文字，可能是掃描檔")
    return tuple(pages)


# 非分頁文件以第一頁證據保存。
def _single_page(text: str) -> tuple[DocumentPageEvidence, ...]:
    return (DocumentPageEvidence(1, text),)


# 擷取 DOCX 的段落與表格儲存格文字。
def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    text = _normalize_text("\n".join([*paragraphs, *table_text]))
    if not text:
        raise ValueError("DOCX 沒有可擷取文字")
    return text


# 從 ODT 壓縮檔中的 content.xml 擷取文字節點。
def _extract_odt_text(content: bytes) -> str:
    try:
        with ZipFile(BytesIO(content)) as archive:
            xml_content = archive.read("content.xml")
    except (BadZipFile, KeyError) as error:
        raise ValueError("ODT 結構無效或缺少 content.xml") from error
    root = ElementTree.fromstring(xml_content)
    text = _normalize_text(" ".join(value for value in root.itertext() if value))
    if not text:
        raise ValueError("ODT 沒有可擷取文字")
    return text


# 壓縮文件文字中的空白與空行。
def _normalize_text(text: str) -> str:
    return " ".join(text.split())

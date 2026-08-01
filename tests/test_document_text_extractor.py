# -*- coding: utf-8 -*-

from io import BytesIO
from types import SimpleNamespace
from zipfile import ZipFile

from docx import Document
import pytest

import src.extractors.document_text_extractor as extractor


# 驗證 DOCX 段落與表格文字都會被擷取。
def test_extract_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("申請資格限電子工程系學生。")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "學業平均八十分以上。"
    stream = BytesIO()
    document.save(stream)

    text = extractor.extract_document_text(
        stream.getvalue(),
        extractor.DOCX_MIME,
        "https://example.com/rules.docx",
        max_pdf_pages=10,
    )

    assert "電子工程系" in text
    assert "學業平均八十分以上" in text


# 驗證 PDF 文字會依頁數上限擷取。
def test_extract_pdf_respects_page_limit(monkeypatch: pytest.MonkeyPatch) -> None:
    pages = [SimpleNamespace(extract_text=lambda value=value: value) for value in ("第一頁", "第二頁")]
    monkeypatch.setattr(extractor, "PdfReader", lambda _: SimpleNamespace(pages=pages))

    text = extractor.extract_document_text(
        b"fake-pdf",
        extractor.PDF_MIME,
        "https://example.com/rules.pdf",
        max_pdf_pages=1,
    )

    assert text == "第一頁"


# 驗證掃描型或空白 PDF 沒有文字時採失敗關閉。
def test_empty_pdf_text_is_rejected(monkeypatch: pytest.MonkeyPatch) -> None:
    pages = [SimpleNamespace(extract_text=lambda: "")]
    monkeypatch.setattr(extractor, "PdfReader", lambda _: SimpleNamespace(pages=pages))

    with pytest.raises(ValueError, match="沒有可擷取文字"):
        extractor.extract_document_text(
            b"fake-pdf",
            extractor.PDF_MIME,
            "https://example.com/scan.pdf",
            max_pdf_pages=10,
        )


# 驗證舊版 DOC 會被辨識並回報明確無法安全解析，不得靜默漏掉。
def test_legacy_doc_is_discovered_but_fails_closed() -> None:
    kind = extractor.detect_document_kind(
        "application/msword",
        "https://example.com/form.doc",
    )

    assert kind == "doc"
    with pytest.raises(ValueError, match="舊式 DOC 已發現"):
        extractor.extract_document_text(
            b"legacy-doc",
            "application/msword",
            "https://example.com/form.doc",
            max_pdf_pages=10,
        )


# 驗證 ODT 辦法可在不執行巨集的情況下擷取文字。
def test_extract_odt_text() -> None:
    stream = BytesIO()
    with ZipFile(stream, "w") as archive:
        archive.writestr(
            "content.xml",
            "<document><p>申請資格限大專院校學生。</p><p>平均八十分以上。</p></document>",
        )

    text = extractor.extract_document_text(
        stream.getvalue(),
        extractor.ODT_MIME,
        "https://example.com/rules.odt",
        max_pdf_pages=10,
    )

    assert "大專院校學生" in text
    assert "平均八十分以上" in text
